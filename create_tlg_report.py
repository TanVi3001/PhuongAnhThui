# -*- coding: utf-8 -*-
"""
Tạo bảng phân tích chỉ tiêu tài chính TLG 2023-2025.

Cách dùng nhanh trên Windows:
    py create_tlg_report_fixed.py

Điều kiện:
    - Đặt file .py này cùng thư mục với file Word mẫu "làmm dở.docx" hoặc "làm dở.docx".
    - Cài thư viện: py -m pip install python-docx

Kết quả xuất ra cùng thư mục với file .py:
    - TLG_bang_phan_tich_markdown.md
    - TLG_lam_do_da_dien_day_du.docx
"""

from pathlib import Path
import argparse
import sys

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NO_DATA = 'Không có dữ liệu trong báo cáo'
OUTPUT_DOCX_NAME = 'TLG_lam_do_da_dien_day_du.docx'
OUTPUT_MD_NAME = 'TLG_bang_phan_tich_markdown.md'


def script_dir() -> Path:
    """Thư mục chứa file .py. Khi chạy local sẽ là ví dụ E:\\PhuongAnh."""
    return Path(__file__).resolve().parent


def set_cell_text(cell, text, bold=False, size=7, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = 'Times New Roman'
    if r._element.rPr is not None:
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_table_borders(table):
    """Kẻ viền bảng để mở bằng Word nhìn rõ hơn."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'w:' + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')


def format_vnd(x):
    if x is None:
        return NO_DATA
    if x == 0:
        return '0'
    s = f"{abs(int(round(x))):,}".replace(',', '.')
    return f"({s})" if x < 0 else s


def format_pct(x):
    if x is None:
        return 'Không tính được'
    return f"{x:.2f}".replace('.', ',') + '%'


def is_num(x):
    return isinstance(x, (int, float))


def calc_delta(a, b):
    if not is_num(a) or not is_num(b):
        return None
    return b - a


def calc_pct(a, b):
    if not is_num(a) or not is_num(b) or a == 0:
        return None
    return (b - a) / a * 100

rows = [
# Balance sheet - assets
('A. BẢNG CÂN ĐỐI KẾ TOÁN HỢP NHẤT - TÀI SẢN', None, None, None, True),
('TÀI SẢN NGẮN HẠN', 2074455057354, 2670706305975, 2748352806718, False),
('Tiền và các khoản tương đương tiền', 243232641902, 700074273053, 487028475726, False),
('Tiền', 243232641902, 334874273053, 309028475726, False),
('Các khoản tương đương tiền', 0, 365200000000, 178000000000, False),
('Đầu tư tài chính ngắn hạn', 447342000000, 360900000000, 423728671233, False),
('Đầu tư nắm giữ đến ngày đáo hạn', 447342000000, 360900000000, 423728671233, False),
('Các khoản phải thu ngắn hạn', 438809370785, 649050923753, 801285542804, False),
('Phải thu ngắn hạn của khách hàng', 390722017248, 612821383076, 771186802967, False),
('Trả trước cho người bán ngắn hạn', 12800019323, 14624171905, 10359462173, False),
('Phải thu về cho vay ngắn hạn', 16000000000, 16000000000, 16000000000, False),
('Phải thu ngắn hạn khác', 30011672235, 12303121563, 14064261122, False),
('Dự phòng phải thu ngắn hạn khó đòi', -10775506081, -6697752791, -10324983458, False),
('Tài sản thiếu chờ xử lý', 51168060, 0, None, False),
('Hàng tồn kho', 831899830043, 784896481994, 814530054475, False),
('Hàng tồn kho (giá gốc)', 855426147198, 833963734982, 868126394694, False),
('Dự phòng giảm giá hàng tồn kho', -23726317155, -49067252988, -53596340219, False),
('Tài sản ngắn hạn khác', 113371161624, 175784627175, 221780062480, False),
('Chi phí trả trước ngắn hạn', 15759857109, 23135505351, 17359481374, False),
('Thuế GTGT được khấu trừ', 93018891136, 143451017773, 203737172831, False),
('Thuế và các khoản phải thu Nhà nước', 4592403379, 9198104051, 683408275, False),
('TÀI SẢN DÀI HẠN', 733993121467, 688855057514, 820392113022, False),
('Các khoản phải thu dài hạn', 12978175740, 14717155781, 13073736612, False),
('Phải thu dài hạn khác', 12978175740, 14717155781, 13073736612, False),
('Tài sản cố định', 575586339381, 528291326958, 523038702112, False),
('Tài sản cố định hữu hình', 549978973266, 504876115604, 509606461579, False),
('Nguyên giá TSCĐ hữu hình', 1309168718614, 1341014256759, 1382068596773, False),
('Giá trị khấu hao lũy kế TSCĐ hữu hình', -759189745348, -836138141155, -872462135194, False),
('Tài sản cố định vô hình', 25607366115, 23415211354, 13432240533, False),
('Nguyên giá TSCĐ vô hình', 79794421651, 81114342736, 73387110236, False),
('Giá trị hao mòn lũy kế TSCĐ vô hình', -54187055536, -57699131382, -59954869703, False),
('Tài sản dở dang dài hạn', 7742307132, 20391542457, 30146867330, False),
('Chi phí xây dựng cơ bản dở dang', 7742307132, 20391542457, 30146867330, False),
('Đầu tư tài chính dài hạn', 62833024000, 64214208000, 182033921343, False),
('Đầu tư vào công ty liên kết', 36000000000, 37200000000, 179913921343, False),
('Đầu tư góp vốn vào đơn vị khác', 30685000000, 30685000000, 5685000000, False),
('Dự phòng đầu tư tài chính dài hạn', -3851976000, -3670792000, -3565000000, False),
('Tài sản dài hạn khác', 74853275214, 61240824318, 72098885625, False),
('Chi phí trả trước dài hạn', 49846542512, 48699473781, 56723234964, False),
('Tài sản thuế thu nhập hoãn lại', 25006732702, 12541350537, 15375650661, False),
('TỔNG TÀI SẢN', 2808448178821, 3359561363489, 3568744919740, False),
# liabilities equity
('B. BẢNG CÂN ĐỐI KẾ TOÁN HỢP NHẤT - NGUỒN VỐN', None, None, None, True),
('NỢ PHẢI TRẢ', 714468096169, 1012333480318, 1046056115990, False),
('Nợ ngắn hạn', 659537277010, 985667030318, 1011881334190, False),
('Phải trả người bán ngắn hạn', 193822557317, 260159397874, 270211290448, False),
('Người mua trả tiền trước ngắn hạn', 3780419107, 8135085608, 17490859596, False),
('Thuế và các khoản phải nộp Nhà nước', 31403758947, 48759143419, 46861033289, False),
('Phải trả người lao động', 40388971499, 43977764490, 42313811530, False),
('Chi phí phải trả ngắn hạn', 91872424289, 81163548557, 142555473140, False),
('Phải trả ngắn hạn khác', 6419625092, 7621951212, 6005004701, False),
('Vay ngắn hạn', 250142393852, 486603154331, 439255830659, False),
('Quỹ khen thưởng, phúc lợi', 41727126827, 49246984827, 47188030827, False),
('Nợ dài hạn', 54930819159, 26666450000, 34174781800, False),
('Vay dài hạn', 33498604659, 6195560000, 6195560000, False),
('Thuế thu nhập hoãn lại phải trả', None, None, 8059346717, False),
('Dự phòng phải trả dài hạn', 21432214500, 20470890000, 19919875083, False),
('VỐN CHỦ SỞ HỮU', 2093980082652, 2347227883171, 2522688803750, False),
('Vốn chủ sở hữu', 2093980082652, 2347227883171, 2522688803750, False),
('Vốn góp của chủ sở hữu', 785944530000, 864535750000, 965283400000, False),
('Cổ phiếu phổ thông có quyền biểu quyết', 785944530000, 864535750000, 965283400000, False),
('Thặng dư vốn cổ phần', 361633483771, 361633483771, 361633483771, False),
('Chênh lệch tỷ giá hối đoái', None, 2019934276, 3024877077, False),
('Quỹ đầu tư phát triển', 261896462556, 261896462556, 261896462556, False),
('Lợi nhuận sau thuế chưa phân phối', 683572010385, 857949709802, 932446514528, False),
('LNST chưa phân phối lũy kế các năm trước', 451323152734, 489876419155, 593725202566, False),
('LNST chưa phân phối của năm nay', 232248857651, 368073290647, 338721311962, False),
('Lợi ích cổ đông không kiểm soát', 346564704, -807457234, -1595934182, False),
('TỔNG NGUỒN VỐN', 2808448178821, 3359561363489, 3568744919740, False),
# Income statement
('C. BÁO CÁO KẾT QUẢ HOẠT ĐỘNG KINH DOANH HỢP NHẤT', None, None, None, True),
('Doanh thu bán hàng và cung cấp dịch vụ', 3496671800535, 3772752571792, 4186212267474, False),
('Các khoản giảm trừ doanh thu', -34798584339, -14167235011, -12282128866, False),
('Doanh thu thuần về bán hàng và cung cấp dịch vụ', 3461873216196, 3758585336781, 4173930138608, False),
('Giá vốn hàng bán và dịch vụ cung cấp', -1949334544727, -2083184660839, -2105352140485, False),
('Lợi nhuận gộp về bán hàng và cung cấp dịch vụ', 1512538671469, 1675400675942, 2068577998123, False),
('Doanh thu hoạt động tài chính', 48462393200, 58322770180, 56923912854, False),
('Chi phí tài chính', -25638774592, -28050598215, -35647806696, False),
('Trong đó: Chi phí lãi vay', -17849057542, -13516022950, -21879953936, False),
('Phần lãi/(lỗ) trong công ty liên kết', -4000000000, 1200000000, 2197434843, False),
('Chi phí bán hàng', -740549101955, -769681526258, -1238932069909, False),
('Chi phí quản lý doanh nghiệp', -346843180321, -357185352088, -330162564841, False),
('Lợi nhuận thuần từ hoạt động kinh doanh', 443970007801, 580025969561, 522956904374, False),
('Thu nhập khác', 9198976721, 10939784326, 55479914862, False),
('Chi phí khác', -1208328608, -4105148101, -6455033835, False),
('Lợi nhuận khác', 7990648113, 6834636225, 49024881027, False),
('Tổng lợi nhuận kế toán trước thuế', 451960655914, 586860605786, 571981785401, False),
('Chi phí thuế TNDN hiện hành', -93999785980, -114197366142, -121700049530, False),
('Chi phí thuế TNDN hoãn lại', -1788597353, -12465382165, -5225046593, False),
('Lợi nhuận sau thuế TNDN', 356174272581, 460197857479, 445056689278, False),
('Lợi nhuận sau thuế phân bổ cho cổ đông của Công ty', 358940537151, 461667743647, 446474886962, False),
('Lợi nhuận sau thuế phân bổ cho cổ đông không kiểm soát', -2766264570, -1469886168, -1418197684, False),
('Lãi cơ bản trên cổ phiếu', 3770, 4808, 3993, False),
('Lãi suy giảm trên cổ phiếu', 3770, 4808, 3993, False),
]




def build_computed_rows():
    computed = []
    for name, y2023, y2024, y2025, section in rows:
        if section:
            computed.append([name, '', '', '', '', '', '', ''])
            continue

        d24 = calc_delta(y2023, y2024)
        p24 = calc_pct(y2023, y2024)
        d25 = calc_delta(y2024, y2025)
        p25 = calc_pct(y2024, y2025)

        computed.append([
            name,
            format_vnd(y2023),
            format_vnd(y2024),
            format_vnd(y2025),
            format_vnd(d24) if d24 is not None else 'Không tính được',
            format_pct(p24),
            format_vnd(d25) if d25 is not None else 'Không tính được',
            format_pct(p25),
        ])
    return computed


def build_markdown(computed):
    headers = [
        'Chỉ tiêu', '2023', '2024', '2025',
        '2024 so với 2023 - Số tiền VND',
        '2024 so với 2023 - Phần trăm (%)',
        '2025 so với 2024 - Số tiền VND',
        '2025 so với 2024 - Phần trăm (%)',
    ]
    md = '| ' + ' | '.join(headers) + ' |\n'
    md += '| ' + ' | '.join(['---'] * len(headers)) + ' |\n'
    for row in computed:
        md += '| ' + ' | '.join(str(x).replace('\n', '<br>') for x in row) + ' |\n'
    return md


def find_template_docx(base_dir: Path, explicit_path: str | None = None) -> Path | None:
    """Tìm file Word mẫu trong cùng thư mục, không dùng /mnt/data cố định nữa."""
    if explicit_path:
        p = Path(explicit_path)
        if not p.is_absolute():
            p = base_dir / p
        if p.exists():
            return p
        raise FileNotFoundError(f'Không tìm thấy file Word mẫu: {p}')

    candidate_names = [
        'làmm dở.docx',
        'làm dở.docx',
        'lam do.docx',
        'lam_do.docx',
        'lam_do_mau.docx',
    ]
    for name in candidate_names:
        p = base_dir / name
        if p.exists():
            return p

    # Nếu không đúng tên, tự lấy file .docx đầu tiên không phải file output/temp.
    for p in sorted(base_dir.glob('*.docx')):
        if p.name.startswith('~$'):
            continue
        if p.name == OUTPUT_DOCX_NAME:
            continue
        return p
    return None


def prepare_document(template_path: Path | None) -> Document:
    if template_path is not None:
        return Document(str(template_path))
    # Fallback: không có file mẫu thì tạo file Word mới thay vì crash.
    return Document()


def ensure_landscape(doc: Document):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.35)
    sec.bottom_margin = Inches(0.35)
    sec.left_margin = Inches(0.30)
    sec.right_margin = Inches(0.30)


def add_title(doc: Document):
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]
    else:
        p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BẢNG PHÂN TÍCH CHỈ TIÊU TÀI CHÍNH HỢP NHẤT TLG 2023-2025')
    run.bold = True
    run.font.name = 'Times New Roman'
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(11)


def get_or_create_table(doc: Document):
    # Dùng bảng mẫu nếu có ít nhất 8 cột; nếu không thì tạo bảng mới.
    if doc.tables and len(doc.tables[0].columns) >= 8:
        table = doc.tables[0]
        while len(table.rows) < 2:
            table.add_row()
        while len(table.rows) > 2:
            table._tbl.remove(table.rows[-1]._tr)
        return table

    table = doc.add_table(rows=2, cols=8)
    return table


def fill_table(table, computed):
    hdr1 = [
        'Chỉ tiêu', 'Năm', 'Năm', 'Năm',
        'Năm 2024 so với năm 2023', 'Năm 2024 so với năm 2023',
        'Năm 2025 so với năm 2024', 'Năm 2025 so với năm 2024',
    ]
    hdr2 = [
        'Chỉ tiêu', '2023', '2024', '2025',
        'Số tiền VND', 'Phần trăm (%)', 'Số tiền VND', 'Phần trăm (%)',
    ]

    for j, txt in enumerate(hdr1):
        set_cell_text(table.rows[0].cells[j], txt, True, 7)
        set_shading(table.rows[0].cells[j], 'D9EAF7')
    for j, txt in enumerate(hdr2):
        set_cell_text(table.rows[1].cells[j], txt, True, 7)
        set_shading(table.rows[1].cells[j], 'D9EAF7')

    for row_data in computed:
        row = table.add_row()
        is_section = (row_data[1] == '')
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(row.cells[j], val, bold=is_section, size=6 if not is_section else 7, align=align)
            if is_section:
                set_shading(row.cells[j], 'E2F0D9')

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    widths = [2300, 1350, 1350, 1350, 1550, 1150, 1550, 1150]
    for row in table.rows:
        for idx, cell in enumerate(row.cells[:8]):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[idx]))
            tcW.set(qn('w:type'), 'dxa')


def main():
    parser = argparse.ArgumentParser(description='Tạo file Word và Markdown phân tích chỉ tiêu tài chính TLG.')
    parser.add_argument('--template', default=None, help='Đường dẫn file Word mẫu. Mặc định tự tìm trong cùng thư mục.')
    parser.add_argument('--out-docx', default=OUTPUT_DOCX_NAME, help='Tên/đường dẫn file Word xuất ra.')
    parser.add_argument('--out-md', default=OUTPUT_MD_NAME, help='Tên/đường dẫn file Markdown xuất ra.')
    args = parser.parse_args()

    base_dir = script_dir()
    computed = build_computed_rows()

    out_md = Path(args.out_md)
    if not out_md.is_absolute():
        out_md = base_dir / out_md
    out_docx = Path(args.out_docx)
    if not out_docx.is_absolute():
        out_docx = base_dir / out_docx

    md = build_markdown(computed)
    out_md.write_text(md, encoding='utf-8')

    template_path = find_template_docx(base_dir, args.template)
    doc = prepare_document(template_path)
    ensure_landscape(doc)
    add_title(doc)
    table = get_or_create_table(doc)
    fill_table(table, computed)
    doc.save(str(out_docx))

    print('DONE - Đã tạo file:')
    print(f'  Word    : {out_docx}')
    print(f'  Markdown: {out_md}')
    if template_path:
        print(f'  File mẫu: {template_path}')
    else:
        print('  File mẫu: không tìm thấy, đã tạo Word mới.')


if __name__ == '__main__':
    try:
        main()
    except Exception as exc:
        print('\nLỖI:', exc, file=sys.stderr)
        print('\nGợi ý sửa:', file=sys.stderr)
        print('- Đặt file .py cùng thư mục với file Word mẫu "làmm dở.docx" hoặc dùng --template "duong_dan_file.docx".', file=sys.stderr)
        print('- Cài thư viện: py -m pip install python-docx', file=sys.stderr)
        sys.exit(1)
